import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue professional headings
    return heading

def create_comprehensive_report():
    doc = docx.Document()
    
    # Title Page
    title = doc.add_heading('PhishGuard: Comprehensive Automated Phishing Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Project Overview & Technical Architecture Report').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Author: Ahmed Abdulaal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Email: ahmed.abdulaal04@gmail.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: May 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', level=1)
    doc.add_paragraph(
        'PhishGuard is an advanced, end-to-end web application and background daemon designed to '
        'automatically detect, classify, and mitigate phishing email attacks in real-time. '
        'By combining a traditional Rule-Based heuristic engine with a state-of-the-art Deep Learning '
        'Neural Network trained on the SpamAssassin and other comprehensive datasets, PhishGuard achieves '
        'high precision and recall. The system features a responsive dark-themed dashboard for manual analysis, '
        'historical threat logging, and real-time inbox monitoring via IMAP.'
    )

    # 2. System Architecture
    add_heading(doc, '2. System Architecture', level=1)
    doc.add_paragraph(
        'The architecture of the PhishGuard ecosystem is composed of three primary interconnected modules:'
    )
    doc.add_heading('2.1 Frontend User Interface', level=2)
    doc.add_paragraph(
        'The web interface is built using modern HTML5, CSS3, and JavaScript, running on top of a Flask backend. '
        'It features a sleek, dark-themed UI that includes:\n'
        '- Home Dashboard: For manual email analysis via drag-and-drop or text input.\n'
        '- Analytics Page: Providing comparative metrics (Accuracy, F1-Score, Confusion Matrices) between detection models.\n'
        '- Threat Intelligence Logs: A historical grid view of previously detected phishing attempts with calculated risk scores.'
    )
    
    p_img1 = doc.add_paragraph()
    r_img1 = p_img1.add_run('[ PLACEHOLDER: Insert "Screenshot 2026-05-09 092040.png" (Main Dashboard) here ]')
    r_img1.font.color.rgb = RGBColor(255, 0, 0)
    r_img1.font.bold = True

    doc.add_heading('2.2 Backend & API (Flask)', level=2)
    doc.add_paragraph(
        'The backend serves as the core orchestrator, built on Python Flask. It exposes RESTful API endpoints for:\n'
        '- /analyze: Receives email body text, processes it through TF-IDF and Keras Text Vectorizers, and routes it to the AI models.\n'
        '- /inbox-stats: A polling endpoint that provides real-time counts of scanned and detected emails to the frontend.\n'
        'The backend ensures secure and efficient memory management by lazy-loading the Keras (.keras) models and Scikit-Learn vectorizers into memory upon the first request.'
    )

    doc.add_heading('2.3 Background IMAP Daemon', level=2)
    doc.add_paragraph(
        'A headless Python daemon (imap_monitor.py) runs continuously in the background. It utilizes the IMAP protocol over SSL (Port 993) '
        'to establish a secure connection to the user\'s inbox (e.g., ahmed.abdulaal04@gmail.com). It polls the INBOX and Spam folders '
        'every 15 seconds, extracting the subject, sender, and payload of unread emails. It then dispatches this data to the local Flask API '
        'for instant classification, triggering a Windows Desktop Notification if a threat is detected.'
    )

    # 3. Detection Models
    add_heading(doc, '3. Dual-Engine Detection Mechanisms', level=1)
    doc.add_paragraph(
        'To ensure robustness against zero-day phishing attacks, PhishGuard employs two distinct methodologies:'
    )
    
    doc.add_heading('3.1 Rule-Based Heuristic Engine', level=2)
    doc.add_paragraph(
        'This engine uses regular expressions and string-matching algorithms to detect common phishing patterns, such as:\n'
        '- Urgency Keywords: "Verify your account", "Suspended", "Immediate action required".\n'
        '- Suspicious Links: Use of IP addresses in URLs, mismatching anchor tags, and URL shorteners.\n'
        '- Sender Discrepancies: Mismatches between the display name and the actual SMTP envelope sender.'
    )

    doc.add_heading('3.2 Deep Learning Neural Network', level=2)
    doc.add_paragraph(
        'The primary classification engine is a Keras-based Neural Network. The model ingests text that has been pre-processed '
        'by a custom Text Vectorization layer (removing stop words, stemming, and tokenization). The architecture consists of:\n'
        '- An Embedding Layer to capture semantic relationships between words.\n'
        '- Dense Layers with ReLU activation and Dropout for regularization to prevent overfitting.\n'
        '- A final Sigmoid output layer to provide a probabilistic Risk Score (0% to 100%).'
    )
    
    p_img2 = doc.add_paragraph()
    r_img2 = p_img2.add_run('[ PLACEHOLDER: Insert "Screenshot 2026-05-09 092055.png" (Analytics Dashboard) here ]')
    r_img2.font.color.rgb = RGBColor(255, 0, 0)
    r_img2.font.bold = True

    # 4. Results & Analytics
    add_heading(doc, '4. Model Performance & Analytics', level=1)
    doc.add_paragraph(
        'The Analytics dashboard provides a clear, data-driven comparison of the two models based on historical scanning data. '
        'As evidenced by the testing phase, the Neural Network significantly outperforms the Rule-Based approach:\n'
        '- Neural Network Accuracy: ~92.33% with an F1-Score of 90.71%.\n'
        '- Rule-Based Accuracy: ~61.27% (High Precision, but very low Recall due to strict rules).\n'
        'The system also generates dynamic Confusion Matrices to track True Positives, False Positives, True Negatives, and False Negatives, '
        'allowing for continuous model fine-tuning.'
    )

    # 5. Threat Intelligence Logging
    add_heading(doc, '5. Threat Intelligence Logging', level=1)
    doc.add_paragraph(
        'All detected phishing attempts are permanently logged in the Threat Intelligence database. Security analysts can '
        'review these logs via the dedicated Threat Logs page. Each log entry acts as a historical record detailing:\n'
        '- The precise timestamp of the attack.\n'
        '- The target email subject and the forged sender identity.\n'
        '- The exact calculated Risk Percentage.\n'
        '- The specific model (Rule-Based or Neural Network) that successfully flagged the email.'
    )
    
    p_img3 = doc.add_paragraph()
    r_img3 = p_img3.add_run('[ PLACEHOLDER: Insert "Screenshot 2026-05-09 092109.png" (Threat Logs) here ]')
    r_img3.font.color.rgb = RGBColor(255, 0, 0)
    r_img3.font.bold = True

    # 6. Conclusion
    add_heading(doc, '6. Conclusion', level=1)
    doc.add_paragraph(
        'PhishGuard successfully bridges the gap between passive email filtering and proactive, real-time user alerting. '
        'By leveraging advanced Deep Learning alongside traditional heuristics, it provides a highly reliable, '
        'explainable, and fast defense mechanism against sophisticated social engineering and phishing attacks.'
    )

    doc.save('PhishGuard_Comprehensive_Report.docx')
    print("Report generated successfully as 'PhishGuard_Comprehensive_Report.docx'.")

if __name__ == "__main__":
    create_comprehensive_report()
